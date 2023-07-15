from flask import Flask, render_template, request, redirect, url_for, session
from flask_sqlalchemy import SQLAlchemy
from datetime import datetime
from openpyxl import Workbook
from docx import Document
import os, secrets

app = Flask(__name__)
app.secret_key = 'qwerty123456'
# app.secret_key = secrets.token_hex(16)
# app.secret_key = os.environ.get('SECRET_KEY')
app.config['SQLALCHEMY_DATABASE_URI'] = 'sqlite:///blog.db'
db = SQLAlchemy(app)

# Define database models
class User(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    username = db.Column(db.String(50), unique=True, nullable=False)
    password = db.Column(db.String(100), nullable=False)

class Post(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    title = db.Column(db.String(100), nullable=False)
    content = db.Column(db.Text, nullable=False)
    author_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    author = db.relationship('User', backref=db.backref('posts', lazy=True))
    created_at = db.Column(db.DateTime, default=datetime.utcnow)

    def save(self):
        db.session.add(self)
        db.session.commit()

    def update(self):
        db.session.commit()

    def delete(self):
        db.session.delete(self)
        db.session.commit()

@app.route('/', methods=['GET'])
def index():
    search_query = request.args.get('search', '')
    if search_query:
        posts = Post.query.filter(Post.title.contains(search_query)).order_by(Post.created_at.desc()).all()
    else:
        posts = Post.query.order_by(Post.created_at.desc()).all()
    return render_template('index.html', posts=posts, search_query=search_query)

@app.route('/register', methods=['GET', 'POST'])
def register():
    if request.method == 'POST':
        username = request.form['username']
        password = request.form['password']
        confirm_password = request.form['confirm_password']
        if password != confirm_password:
            error = 'Passwords do not match.'
            return render_template('register.html', error=error)
        existing_user = User.query.filter_by(username=username).first()
        if existing_user:
            error = 'Username already exists.'
            return render_template('register.html', error=error)
        new_user = User(username=username, password=password)
        db.session.add(new_user)
        db.session.commit()
        return redirect(url_for('login'))
    return render_template('register.html')


@app.route('/login', methods=['GET', 'POST'])
def login():
    if request.method == 'POST':
        username = request.form['username']
        password = request.form['password']
        # Validate user credentials
        user = User.query.filter_by(username=username).first()
        if user and user.password == password:
            # Set user session or generate authentication token
            session['user_id'] = user.id
            return redirect(url_for('index'))
        else:
            error = 'Invalid username or password'
            return render_template('login.html', error=error)
    return render_template('login.html')

@app.route('/logout')
def logout():
    # Clear user session
    session.clear()
    return redirect(url_for('index'))

@app.route('/posts/new', methods=['GET', 'POST'])
def create_post():
    if 'user_id' not in session:
        return redirect(url_for('login'))
    if request.method == 'POST':
        title = request.form['title']
        content = request.form['content']
        author_id = session['user_id']
        new_post = Post(title=title, content=content, author_id=author_id)
        db.session.add(new_post)
        db.session.commit()
        return redirect(url_for('index'))
    return render_template('create_post.html')

@app.route('/posts/<int:post_id>/edit', methods=['GET', 'POST'])
def edit_post(post_id):
    if 'user_id' not in session:
        return redirect(url_for('login'))
    post = Post.query.get(post_id)
    if not post:
        return redirect(url_for('index'))
    if request.method == 'POST':
        title = request.form['title']
        content = request.form['content']
        post.title = title
        post.content = content
        db.session.commit()
        return redirect(url_for('index'))
    return render_template('edit_post.html', post=post)

@app.route('/posts/<int:post_id>/delete', methods=['POST'])
def delete_post(post_id):
    if 'user_id' not in session:
        return redirect(url_for('login'))
    post = Post.query.get_or_404(post_id)
    if post.author_id != session['user_id']:
        return redirect(url_for('index'))
    db.session.delete(post)
    db.session.commit()
    return redirect(url_for('index'))

@app.route('/posts/export/excel')
def export_to_excel():
    posts = Post.query.all()

    workbook = Workbook()
    worksheet = workbook.active

    worksheet.append(["Title", "Content", "Author", "Created At"])

    for post in posts:
        worksheet.append([post.title, post.content, post.author.username, post.created_at])

    workbook.save('blog_posts.xlsx')

    return redirect(url_for('index'))

@app.route('/posts/export/word')
def export_to_word():
    posts = Post.query.all()

    document = Document()

    document.add_heading('Blog Posts', level=1)

    for post in posts:
        document.add_heading(post.title, level=2)
        document.add_paragraph(f"By: {post.author.username}")
        document.add_paragraph(post.content)
        document.add_paragraph(f"Created At: {post.created_at}")

    document.save('blog_posts.docx')

    return redirect(url_for('index'))

if __name__ == '__main__':
    with app.app_context():
        db.create_all()
    app.run(debug=True)